"""DOCX extraction with graceful dependency handling."""

from __future__ import annotations

from pathlib import Path

from src.ingestion.extractors.common import BaseExtractor, ExtractionResult


class DOCXExtractor(BaseExtractor):
    """Extract paragraph and table text from DOCX documents."""

    async def can_extract(self, file_path: Path) -> bool:
        """Return whether the file is a DOCX document."""

        return file_path.suffix.lower() == ".docx"

    async def extract(self, file_path: Path, doc_id: str, folder_id: str) -> ExtractionResult:
        """Extract text content from a DOCX file."""

        document = self._create_document_metadata(file_path, doc_id, folder_id)
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            return ExtractionResult(
                document=document,
                text_content="",
                success=False,
                error_message=f"python-docx is required for DOCX support: {exc}",
            )

        try:
            doc = DocxDocument(file_path)
            paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
            table_blocks: list[str] = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    rows.append(" | ".join(cell.text.strip() for cell in row.cells))
                if rows:
                    table_blocks.append("\n".join(rows))

            document.has_tables = bool(table_blocks)
            document.has_images = any("image" in rel.target_ref.lower() for rel in doc.part.rels.values())
            text_content = "\n\n".join(paragraphs + table_blocks)

            return ExtractionResult(document=document, text_content=text_content, success=True)
        except Exception as exc:
            return ExtractionResult(
                document=document,
                text_content="",
                success=False,
                error_message=str(exc),
            )
