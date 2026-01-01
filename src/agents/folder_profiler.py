import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
import asyncio
from datetime import datetime

from src.agents.base import Agent, AgentResult
from src.core.orchestration.state import RunState
from src.core.policies.folder_policy import (
    FolderPolicy, DocumentType, RetrievalStrategy,
    AuthorityLevel
)
from src.llm.client import LLMClient
from src.utils.fileio import list_files
from src.utils.ids import generate_folder_id

logger = logging.getLogger(__name__)


class FolderScanResult(BaseModel):
    """Structured result from folder scanning"""
    all_files: List[Path]
    text_files: List[Path]
    pdf_files: List[Path]
    docx_files: List[Path]
    xlsx_files: List[Path]
    image_files: List[Path]
    total_count: int
    has_tables: bool
    has_images: bool
    total_size_bytes: int
    
    class Config:
        arbitrary_types_allowed = True


class ImageAnalysisResult(BaseModel):
    """Result from image requirement analysis"""
    image_processing_required: bool
    has_standalone_images: bool
    has_scanned_pdfs: bool
    has_embedded_pdf_images: bool
    has_docx_images: bool
    estimated_image_count: int
    detected_image_types: List[str]  # ["charts", "forms", "screenshots", "diagrams"]
    requires_ocr: bool
    requires_caption: bool
    may_contain_pii: bool


class FolderProfileResult(AgentResult):
    """Result from folder profiling"""
    policy: FolderPolicy


class FolderProfilerAgent(Agent[FolderProfileResult]):
    """
    Analyzes a document folder and generates a comprehensive FolderPolicy.
    
    This agent is critical to the system's intelligence - it determines:
    - What types of documents exist
    - How they should be retrieved (dense/sparse/hybrid)
    - What image processing is needed
    - What entities and topics to expect
    - Authority and governance rules
    
    The generated policy guides all downstream ingestion and retrieval.
    """
    
    def __init__(
        self,
        llm_client: LLMClient,
        deep_scan: bool = True,
        max_samples: int = 5,
        scan_images: bool = True
    ):
        super().__init__(llm_client)
        self.deep_scan = deep_scan  # Whether to do PDF content inspection
        self.max_samples = max_samples
        self.scan_images = scan_images
    
    async def _execute(self, state: RunState) -> FolderProfileResult:
        """
        This agent doesn't use RunState - it operates standalone.
        Use profile_folder() directly instead.
        """
        raise NotImplementedError(
            "FolderProfilerAgent should be called via profile_folder() method, "
            "not through the standard agent pipeline."
        )
    
    async def profile_folder(
        self,
        folder_path: Path,
        existing_policy: Optional[FolderPolicy] = None
    ) -> FolderPolicy:
        """
        Main entry point: analyze folder and return a comprehensive policy.
        
        Args:
            folder_path: Path to the folder to profile
            existing_policy: If provided, update this policy rather than creating new
        
        Returns:
            Generated or updated FolderPolicy
        
        Process:
            1. Generate stable folder ID
            2. Scan folder structure (file types, counts, sizes)
            3. Analyze image requirements (OCR, captions, PII risks)
            4. Sample document content for topics/entities
            5. Generate policy via LLM with all context
            6. Validate and enrich with computed values
        """
        start_time = datetime.utcnow()
        
        self.logger.info(
            f"Starting folder profiling: {folder_path}",
            extra={"deep_scan": self.deep_scan, "scan_images": self.scan_images}
        )
        
        if not folder_path.exists():
            raise FileNotFoundError(f"Folder not found: {folder_path}")
        
        if not folder_path.is_dir():
            raise ValueError(f"Path is not a directory: {folder_path}")
        
        # 1. Generate stable folder ID
        folder_id = generate_folder_id(folder_path)
        
        # 2. Scan folder structure
        self.logger.debug("Scanning folder structure...")
        scan_result = await self._scan_folder(folder_path)
        
        self.logger.info(
            f"Scan complete: {scan_result.total_count} files",
            extra={
                "pdfs": len(scan_result.pdf_files),
                "docx": len(scan_result.docx_files),
                "xlsx": len(scan_result.xlsx_files),
                "images": len(scan_result.image_files)
            }
        )
        
        # 3. Analyze image requirements
        self.logger.debug("Analyzing image requirements...")
        image_analysis = await self._analyze_image_requirements(scan_result)
        
        self.logger.info(
            f"Image analysis complete",
            extra={
                "processing_required": image_analysis.image_processing_required,
                "scanned_pdfs": image_analysis.has_scanned_pdfs,
                "estimated_images": image_analysis.estimated_image_count
            }
        )
        
        # 4. Sample documents for content analysis
        self.logger.debug(f"Sampling up to {self.max_samples} documents...")
        content_samples = await self._sample_documents(
            folder_path,
            scan_result,
            max_samples=self.max_samples
        )
        
        self.logger.info(f"Collected {len(content_samples)} content samples")
        
        # 5. Generate policy with LLM
        self.logger.debug("Generating policy with LLM...")
        policy = await self._generate_policy_with_llm(
            folder_path=folder_path,
            folder_id=folder_id,
            scan_result=scan_result,
            image_analysis=image_analysis,
            content_samples=content_samples,
            existing_policy=existing_policy
        )
        
        # 6. Validate and enrich
        policy = self._enrich_policy(
            policy,
            folder_path,
            folder_id,
            scan_result,
            image_analysis
        )
        
        duration = (datetime.utcnow() - start_time).total_seconds()
        
        self.logger.info(
            f"Policy generation complete for {folder_path.name}",
            extra={
                "duration_seconds": duration,
                "document_types": [dt.value for dt in policy.document_types],
                "retrieval_strategy": policy.retrieval_strategy.value,
                "image_processing": policy.image_processing_required,
                "authority": policy.authority_level.value
            }
        )
        
        return policy
    
    async def _scan_folder(self, folder_path: Path) -> FolderScanResult:
        """
        Recursively scan folder and collect file statistics.
        
        Returns structured information about all files, organized by type.
        """
        # Define supported extensions
        EXT_MAP = {
            'text': ['.txt', '.md', '.rst'],
            'pdf': ['.pdf'],
            'docx': ['.docx', '.doc'],
            'xlsx': ['.xlsx', '.xls', '.csv'],
            'image': ['.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.webp', '.gif']
        }
        
        # Initialize collections
        all_files = []
        categorized = {key: [] for key in EXT_MAP.keys()}
        total_size = 0
        
        # Recursively walk folder
        for item in folder_path.rglob('*'):
            # Skip hidden files and directories
            if item.name.startswith('.') or any(part.startswith('.') for part in item.parts):
                continue
            
            if item.is_file():
                all_files.append(item)
                total_size += item.stat().st_size
                
                ext = item.suffix.lower()
                
                # Categorize by extension
                for category, extensions in EXT_MAP.items():
                    if ext in extensions:
                        categorized[category].append(item)
                        break
        
        return FolderScanResult(
            all_files=all_files,
            text_files=categorized['text'],
            pdf_files=categorized['pdf'],
            docx_files=categorized['docx'],
            xlsx_files=categorized['xlsx'],
            image_files=categorized['image'],
            total_count=len(all_files),
            has_tables=len(categorized['xlsx']) > 0,
            has_images=len(categorized['image']) > 0,
            total_size_bytes=total_size
        )
    
    async def _analyze_image_requirements(
        self,
        scan_result: FolderScanResult
    ) -> ImageAnalysisResult:
        """
        Determine what image processing is needed for this folder.
        
        Checks for:
        - Standalone image files
        - Scanned PDFs (pages as images)
        - Embedded PDF images (figures/charts)
        - DOCX embedded images
        - Image content types (forms, charts, diagrams)
        - PII risks
        """
        has_standalone = len(scan_result.image_files) > 0
        has_pdfs = len(scan_result.pdf_files) > 0
        has_docx = len(scan_result.docx_files) > 0
        
        has_scanned_pdfs = False
        has_embedded_pdf_images = False
        has_docx_images = False
        detected_image_types = []
        estimated_image_count = len(scan_result.image_files)
        
        # Analyze PDFs if deep_scan enabled
        if has_pdfs and self.deep_scan and self.scan_images:
            pdf_analysis = await self._analyze_pdfs_for_images(
                scan_result.pdf_files[:10]  # Sample first 10 PDFs
            )
            
            has_scanned_pdfs = pdf_analysis['has_scanned']
            has_embedded_pdf_images = pdf_analysis['has_embedded']
            estimated_image_count += pdf_analysis['estimated_count']
            detected_image_types.extend(pdf_analysis['detected_types'])
        
        # Analyze DOCX if deep_scan enabled
        if has_docx and self.deep_scan and self.scan_images:
            docx_analysis = await self._analyze_docx_for_images(
                scan_result.docx_files[:10]  # Sample first 10 DOCX
            )
            
            has_docx_images = docx_analysis['has_images']
            estimated_image_count += docx_analysis['estimated_count']
        
        # Determine if image processing is required
        image_processing_required = (
            has_standalone or
            has_scanned_pdfs or
            has_embedded_pdf_images or
            has_docx_images
        )
        
        # Decide on OCR and captioning
        requires_ocr = (
            has_scanned_pdfs or  # Scanned PDFs always need OCR
            has_standalone or    # Standalone images might have text
            "forms" in detected_image_types or
            "screenshots" in detected_image_types
        )
        
        requires_caption = (
            has_embedded_pdf_images or  # Charts/figures benefit from captions
            "charts" in detected_image_types or
            "diagrams" in detected_image_types or
            len(scan_result.image_files) > 0
        )
        
        # PII risk heuristics
        may_contain_pii = (
            "forms" in detected_image_types or
            "screenshots" in detected_image_types or
            has_scanned_pdfs  # Scanned forms often have PII
        )
        
        # Deduplicate image types
        detected_image_types = list(set(detected_image_types))
        
        return ImageAnalysisResult(
            image_processing_required=image_processing_required,
            has_standalone_images=has_standalone,
            has_scanned_pdfs=has_scanned_pdfs,
            has_embedded_pdf_images=has_embedded_pdf_images,
            has_docx_images=has_docx_images,
            estimated_image_count=estimated_image_count,
            detected_image_types=detected_image_types,
            requires_ocr=requires_ocr,
            requires_caption=requires_caption,
            may_contain_pii=may_contain_pii
        )
    
    async def _analyze_pdfs_for_images(
        self,
        pdf_files: List[Path]
    ) -> Dict[str, Any]:
        """
        Deep scan PDFs to detect scanned pages and embedded images.
        
        Returns dict with:
            - has_scanned: bool
            - has_embedded: bool
            - estimated_count: int
            - detected_types: List[str]
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            self.logger.warning("PyMuPDF not available, skipping PDF image analysis")
            return {
                'has_scanned': False,
                'has_embedded': False,
                'estimated_count': 0,
                'detected_types': []
            }
        
        has_scanned = False
        has_embedded = False
        total_image_count = 0
        detected_types = set()
        
        for pdf_path in pdf_files:
            try:
                doc = fitz.open(pdf_path)
                
                # Check first 3 pages for scanned content
                for page_num in range(min(3, len(doc))):
                    page = doc[page_num]
                    
                    # Get text blocks and images
                    text_blocks = page.get_text("blocks")
                    images = page.get_images()
                    
                    # Heuristic: if page has few text blocks but has images, likely scanned
                    # OR if page has very little extractable text
                    text_content = page.get_text().strip()
                    
                    if (len(text_blocks) < 3 and len(images) > 0) or len(text_content) < 50:
                        has_scanned = True
                        detected_types.add("scanned_pages")
                    
                    # Count embedded images
                    if len(images) > 0:
                        has_embedded = True
                        total_image_count += len(images)
                        
                        # Try to infer image type from size/aspect ratio
                        for img_ref in images[:3]:  # Sample first 3 images per page
                            xref = img_ref[0]
                            try:
                                base_image = doc.extract_image(xref)
                                width = base_image.get("width", 0)
                                height = base_image.get("height", 0)
                                
                                # Heuristics for image types
                                if width > 0 and height > 0:
                                    aspect_ratio = width / height
                                    
                                    # Wide images might be charts
                                    if 1.5 < aspect_ratio < 3.0:
                                        detected_types.add("charts")
                                    
                                    # Square-ish images might be diagrams
                                    elif 0.8 < aspect_ratio < 1.2:
                                        detected_types.add("diagrams")
                                    
                                    # Tall images might be forms
                                    elif aspect_ratio < 0.7:
                                        detected_types.add("forms")
                            
                            except Exception:
                                pass  # Skip if image extraction fails
                
                doc.close()
            
            except Exception as e:
                self.logger.warning(f"Failed to analyze PDF {pdf_path.name}: {e}")
        
        return {
            'has_scanned': has_scanned,
            'has_embedded': has_embedded,
            'estimated_count': total_image_count,
            'detected_types': list(detected_types)
        }
    
    async def _analyze_docx_for_images(
        self,
        docx_files: List[Path]
    ) -> Dict[str, Any]:
        """
        Check DOCX files for embedded images.
        
        Returns dict with:
            - has_images: bool
            - estimated_count: int
        """
        try:
            from docx import Document
        except ImportError:
            self.logger.warning("python-docx not available, skipping DOCX image analysis")
            return {
                'has_images': False,
                'estimated_count': 0
            }
        
        has_images = False
        total_count = 0
        
        for docx_path in docx_files:
            try:
                doc = Document(docx_path)
                
                # Count image relationships
                image_count = sum(
                    1 for rel in doc.part.rels.values()
                    if "image" in rel.target_ref.lower()
                )
                
                if image_count > 0:
                    has_images = True
                    total_count += image_count
            
            except Exception as e:
                self.logger.warning(f"Failed to analyze DOCX {docx_path.name}: {e}")
        
        return {
            'has_images': has_images,
            'estimated_count': total_count
        }
    
    async def _sample_documents(
        self,
        folder_path: Path,
        scan_result: FolderScanResult,
        max_samples: int = 5
    ) -> List[Dict[str, str]]:
        """
        Extract text samples from documents for LLM content analysis.
        
        Returns list of dicts with 'filename' and 'content' keys.
        Prioritizes variety: tries to sample different file types.
        """
        samples = []
        
        # Create a prioritized list of files to sample
        sample_pool = []
        
        # Prioritize one of each type
        for file_list in [
            scan_result.pdf_files[:2],
            scan_result.docx_files[:2],
            scan_result.text_files[:2],
            scan_result.xlsx_files[:1]
        ]:
            sample_pool.extend(file_list)
        
        # Limit to max_samples
        sample_pool = sample_pool[:max_samples]
        
        for file_path in sample_pool:
            try:
                sample_text = await self._extract_sample_text(file_path)
                
                if sample_text and len(sample_text.strip()) > 50:
                    samples.append({
                        'filename': file_path.name,
                        'content': sample_text
                    })
            
            except Exception as e:
                self.logger.warning(f"Failed to sample {file_path.name}: {e}")
        
        return samples
    
    async def _extract_sample_text(self, file_path: Path) -> str:
        """
        Extract sample text from a single document.
        Returns first ~2000 characters.
        """
        ext = file_path.suffix.lower()
        max_chars = 2000
        
        try:
            if ext in ['.txt', '.md', '.rst']:
                # Plain text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read(max_chars)
            
            elif ext == '.pdf':
                # PDF
                import fitz
                doc = fitz.open(file_path)
                text_parts = []
                char_count = 0
                
                for page_num in range(min(3, len(doc))):
                    page_text = doc[page_num].get_text()
                    text_parts.append(page_text)
                    char_count += len(page_text)
                    
                    if char_count >= max_chars:
                        break
                
                doc.close()
                return "\n".join(text_parts)[:max_chars]
            
            elif ext in ['.docx', '.doc']:
                # Word document
                from docx import Document
                doc = Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs[:30]]
                return "\n".join(paragraphs)[:max_chars]
            
            elif ext in ['.xlsx', '.xls']:
                # Excel - just get sheet names and first few cells
                import openpyxl
                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                sheets_info = []
                
                for sheet_name in wb.sheetnames[:3]:
                    ws = wb[sheet_name]
                    cells = []
                    
                    for row in range(1, min(6, ws.max_row + 1)):
                        row_data = []
                        for col in range(1, min(6, ws.max_column + 1)):
                            cell_value = ws.cell(row, col).value
                            if cell_value:
                                row_data.append(str(cell_value))
                        
                        if row_data:
                            cells.append(" | ".join(row_data))
                    
                    sheets_info.append(f"Sheet: {sheet_name}\n" + "\n".join(cells))
                
                wb.close()
                return "\n\n".join(sheets_info)[:max_chars]
            
            else:
                return ""
        
        except Exception as e:
            self.logger.debug(f"Text extraction failed for {file_path.name}: {e}")
            return ""
    
    async def _generate_policy_with_llm(
        self,
        folder_path: Path,
        folder_id: str,
        scan_result: FolderScanResult,
        image_analysis: ImageAnalysisResult,
        content_samples: List[Dict[str, str]],
        existing_policy: Optional[FolderPolicy] = None
    ) -> FolderPolicy:
        """
        Use LLM to generate a comprehensive folder policy.
        
        Provides all context: file counts, image analysis, content samples.
        LLM determines: doc types, domain, topics, entities, retrieval strategy.
        """
        # Format content samples for prompt
        samples_text = ""
        if content_samples:
            samples_text = "\n\n".join([
                f"**File: {sample['filename']}**\n{sample['content'][:1500]}"
                for sample in content_samples
            ])
        else:
            samples_text = "No content samples available."
        
        # Format image types
        image_types_str = ", ".join(image_analysis.detected_image_types) if image_analysis.detected_image_types else "unknown"
        
        # Build comprehensive prompt
        prompt = self.load_prompt(
            "folder_profiler",
            variables={
                "folder_name": folder_path.name,
                "total_files": scan_result.total_count,
                "pdf_count": len(scan_result.pdf_files),
                "docx_count": len(scan_result.docx_files),
                "xlsx_count": len(scan_result.xlsx_files),
                "text_count": len(scan_result.text_files),
                "image_count": len(scan_result.image_files),
                "total_size_mb": round(scan_result.total_size_bytes / (1024 * 1024), 2),
                "has_tables": "Yes" if scan_result.has_tables else "No",
                "has_images": "Yes" if scan_result.has_images else "No",
                "has_scanned_pdfs": "Yes" if image_analysis.has_scanned_pdfs else "No",
                "has_embedded_images": "Yes" if image_analysis.has_embedded_pdf_images else "No",
                "estimated_image_count": image_analysis.estimated_image_count,
                "image_types": image_types_str,
                "requires_ocr": "Yes" if image_analysis.requires_ocr else "No",
                "requires_caption": "Yes" if image_analysis.requires_caption else "No",
                "may_contain_pii": "Yes" if image_analysis.may_contain_pii else "No",
                "content_samples": samples_text,
                "existing_summary": existing_policy.summary if existing_policy else "N/A",
                "update_mode": "Yes - update existing policy" if existing_policy else "No - create new policy"
            }
        )
        
        # Generate structured policy with LLM
        self.logger.debug("Calling LLM for policy generation...")
        
        policy = self.llm.generate_structured(
            prompt=prompt,
            response_model=FolderPolicy,
            system_prompt="""You are an expert document analyst specializing in information retrieval and knowledge management.

                        Your task is to analyze document folders and generate comprehensive policies that guide:
                        - Document ingestion and processing
                        - Image extraction and OCR requirements
                        - Retrieval strategies (vector vs keyword search)
                        - Expected entities and topics
                        - Governance and citation requirements

                        Be thorough, accurate, and pragmatic in your assessments.""",
            temperature=0.2  # Lower temperature for more consistent categorization
        )
        
        return policy
    
    def _enrich_policy(
        self,
        policy: FolderPolicy,
        folder_path: Path,
        folder_id: str,
        scan_result: FolderScanResult,
        image_analysis: ImageAnalysisResult
    ) -> FolderPolicy:
        """
        Enrich LLM-generated policy with computed/override values.
        
        Ensures consistency between LLM output and actual scanned data.
        """
        # Override identity fields
        policy.folder_id = folder_id
        policy.folder_path = str(folder_path.resolve())
        policy.folder_name = folder_path.name
        
        # Override computed statistics
        policy.total_documents = scan_result.total_count
        policy.has_tables = scan_result.has_tables
        policy.has_images = scan_result.has_images
        
        # Override image processing flags (trust scan over LLM)
        policy.image_processing_required = image_analysis.image_processing_required
        policy.has_scanned_pdfs = image_analysis.has_scanned_pdfs
        policy.has_standalone_images = image_analysis.has_standalone_images
        policy.estimated_image_count = image_analysis.estimated_image_count
        policy.image_types = image_analysis.detected_image_types
        policy.ocr_enabled = image_analysis.requires_ocr
        policy.caption_enabled = image_analysis.requires_caption
        policy.images_may_contain_pii = image_analysis.may_contain_pii
        
        # Set timestamps
        policy.updated_at = datetime.utcnow()
        
        # Validate retrieval strategy makes sense
        if policy.retrieval_strategy == RetrievalStrategy.STRUCTURED and not scan_result.has_tables:
            self.logger.warning(
                "LLM selected STRUCTURED strategy but no tables found, changing to HYBRID"
            )
            policy.retrieval_strategy = RetrievalStrategy.HYBRID
        
        # Ensure keyword_boost is reasonable
        if policy.keyword_boost < 0.1 or policy.keyword_boost > 10.0:
            self.logger.warning(
                f"Unusual keyword_boost value: {policy.keyword_boost}, clamping to [0.5, 2.0]"
            )
            policy.keyword_boost = max(0.5, min(2.0, policy.keyword_boost))
        
        return policy